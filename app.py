# app.py - Modified for Multi-language Support

from flask import Flask, render_template, request, redirect, url_for, session, send_file, abort
import pandas as pd
import sqlite3
from sklearn.model_selection import train_test_split
from sklearn.ensemble import RandomForestClassifier
from sklearn.metrics import classification_report
import json
import io
from datetime import datetime, timedelta
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import docx.shared
import joblib

app = Flask(__name__)
app.secret_key = 'supersecretkey'  # Change this for production use

# === TRANSLATION DICTIONARY ===
TRANSLATIONS = {
    'en': {
        'app_name': 'NeuroDetect',
        'welcome_back': 'Welcome back',
        'dashboard': 'Dashboard',
        'profile': 'Profile',
        'cognitive_tests': 'Cognitive Tests',
        'mri_upload': 'MRI Upload',
        'results': 'Results',
        'history': 'History',
        'logout': 'Logout',
        'notifications': 'Notifications',
        'total_tests': 'Total Tests',
        'risk_level': 'Risk Level',
        'avg_score': 'Avg Score',
        'last_test': 'Last Test',
        'ago': 'ago',
        'quick_actions': 'Quick Actions',
        'start_cognitive_test': 'Start Cognitive Test',
        'cognitive_test_desc': 'Begin your comprehensive cognitive assessment',
        'upload_mri_scan': 'Upload MRI Scan',
        'mri_scan_desc': 'Upload brain scan for AI-powered analysis',
        'view_results': 'View Results',
        'view_results_desc': 'Check your latest assessment results',
        'test_history': 'Test History',
        'test_history_desc': 'View all past assessments and trends',
        'recent_tests': 'Recent Tests',
        'view_all': 'View All',
        'mark_all_read': 'Mark all as read',
        'daily_health_tip': '💡 Daily Health Tip',
        'health_tip_text': 'Regular mental exercises and a balanced diet rich in omega-3 fatty acids can help maintain cognitive function. Try puzzle games, meditation, or learn a new skill today!',
        'low_risk': 'Low Risk',
        'high_risk': 'High Risk',
        'score': 'Score',
        'view_details': 'View Details',
        'patient_name': 'Patient Full Name',
        'age': 'Age',
        'gender': 'Gender',
        'male': 'Male',
        'female': 'Female',
        'select_gender': 'Select Gender',
        'submit': 'Submit',
        'generate_assessment': 'Generate Risk Assessment',
        'comprehensive_assessment': 'Comprehensive Risk Assessment',
        'patient_info': 'Patient Information',
        'medical_history': 'Medical & Family History',
        'lifestyle': 'Lifestyle & Physical Health',
        'download_pdf': 'Download PDF Report',
        'download_docx': 'Download Word Report',
        'prediction_result': 'Prediction Result',
        'no_data': 'No data available',
        'update_profile': 'Update Profile',
        'email': 'Email',
        'save_changes': 'Save Changes',
        'language': 'Language',
        'select_language': 'Select Language',
        'login': 'Login',
        'signup': 'Sign Up',
        'username': 'Username',
        'password': 'Password',
        'start_assessment': 'Start Assessment',
        'monitor_health': 'Monitor your cognitive health journey with our AI-powered insights',
        'back_to_dashboard': 'Back to Dashboard',
        'complete_form_desc': 'Complete the following form to generate an AI-powered Alzheimer\'s risk prediction',
        'patient_info_desc': 'Basic demographic and identification details',
        'medical_history_desc': 'Previous conditions and family medical background',
        'lifestyle_desc': 'Daily habits and physical wellness indicators',
        'vital_signs': 'Vital Signs & Laboratory Results',
        'vital_signs_desc': 'Blood pressure, cholesterol, and metabolic markers',
        'cognitive_assessment': 'Cognitive & Functional Assessment',
        'cognitive_assessment_desc': 'Mental capacity and daily functioning scores',
        'clinical_symptoms': 'Clinical Symptoms & Behavioral Indicators',
        'clinical_symptoms_desc': 'Observed symptoms and behavioral changes',
        'enter_patient_name': 'Enter patient\'s full name',
        'years': 'years',
        'ethnicity': 'Ethnicity',
        'select_ethnicity': 'Select Ethnicity',
        'other': 'Other',
        'asian': 'Asian',
        'black': 'Black',
        'white': 'White',
        'education_level': 'Education Level',
        'select_level': 'Select Level',
        'no_schooling': 'No Schooling',
        'primary_education': 'Primary Education',
        'secondary_education': 'Secondary Education',
        'higher_education': 'Higher Education',
        'family_history_alz': 'Family History of Alzheimer\'s',
        'family_history': 'Family History',
        'yes': 'Yes',
        'no': 'No',
        'diabetes': 'Diabetes',
        'hypertension': 'Hypertension',
        'cardiovascular_disease': 'Cardiovascular Disease',
        'depression': 'Depression',
        'head_injury': 'Head Injury',
        'bmi': 'BMI',
        'physical_activity': 'Physical Activity Level',
        'smoking_status': 'Smoking Status',
        'smoking': 'Smoking',
        'non_smoker': 'Non-Smoker',
        'smoker': 'Smoker',
        'alcohol_consumption': 'Alcohol Consumption',
        'diet_quality': 'Diet Quality Score',
        'sleep_quality': 'Sleep Quality Score',
        'systolic_bp': 'Systolic Blood Pressure',
        'diastolic_bp': 'Diastolic Blood Pressure',
        'total_cholesterol': 'Total Cholesterol',
        'ldl_cholesterol': 'LDL Cholesterol',
        'hdl_cholesterol': 'HDL Cholesterol',
        'triglycerides': 'Triglycerides',
        'mmse_score': 'MMSE Score',
        'mmse': 'MMSE',
        'functional_assessment': 'Functional Assessment Score',
        'adl_score': 'ADL Score',
        'adl': 'ADL',
        'memory_complaints': 'Memory Complaints',
        'behavioral_problems': 'Behavioral Problems',
        'confusion': 'Confusion',
        'disorientation': 'Disorientation',
        'personality_changes': 'Personality Changes',
        'difficulty_completing_tasks': 'Difficulty Completing Tasks',
        'forgetfulness': 'Forgetfulness',
        'ai_powered_platform': 'AI-Powered Early Detection Platform',
        'disclaimer_text': 'This tool is for screening purposes only and does not replace professional medical diagnosis.',
        'age_tooltip': 'Your current age matters because Alzheimer\'s risk increases as you get older, especially after 65.',
        'gender_tooltip': 'Women tend to have slightly higher risk of Alzheimer\'s, partly because they generally live longer than men.',
        'ethnicity_tooltip': 'Some ethnic groups have different risk levels for Alzheimer\'s due to genetic and health factors.',
        'education_tooltip': 'Higher education helps build brain strength, which may protect against memory problems later in life.',
        'family_history_tooltip': 'If your parents or siblings had Alzheimer\'s, you may have a higher chance of developing it too.',
        'diabetes_tooltip': 'High blood sugar from diabetes can damage your brain over time and increase memory problem risks.',
        'hypertension_tooltip': 'High blood pressure can damage blood vessels in your brain, reducing blood flow and affecting memory.',
        'cardiovascular_tooltip': 'Heart problems can affect blood flow to your brain, which may increase the risk of memory issues.',
        'depression_tooltip': 'Long-term depression can affect your brain health and may be connected to higher Alzheimer\'s risk.',
        'head_injury_tooltip': 'Serious head injuries or repeated concussions can increase your risk of developing memory problems later.',
        'bmi_tooltip': 'Your weight compared to height. Being very overweight or underweight can affect brain health.',
        'physical_activity_tooltip': 'How much you exercise. Regular activity keeps your brain healthy. 0 means no exercise, 10 means daily workouts.',
        'smoking_tooltip': 'Smoking harms blood vessels in your brain and increases your risk of memory problems significantly.',
        'alcohol_tooltip': 'How many alcoholic drinks you have per week. Heavy drinking can damage your brain over time.',
        'diet_quality_tooltip': 'How healthy you eat. More fruits, vegetables, and fish help brain health. 0 is poor diet, 10 is excellent.',
        'sleep_quality_tooltip': 'How well you sleep. Good sleep helps clean toxins from your brain. 0 is very poor sleep, 10 is excellent.',
        'systolic_bp_tooltip': 'The top number in your blood pressure reading. High numbers can damage blood vessels in your brain.',
        'diastolic_bp_tooltip': 'The bottom number in your blood pressure reading. It shows pressure when your heart rests between beats.',
        'cholesterol_total_tooltip': 'All the cholesterol in your blood. Very high levels can harm blood vessels feeding your brain.',
        'ldl_tooltip': 'The \'bad\' cholesterol that can clog blood vessels. Lower numbers are better for brain health.',
        'hdl_tooltip': 'The \'good\' cholesterol that helps remove bad cholesterol. Higher numbers protect your brain.',
        'triglycerides_tooltip': 'A type of fat in your blood. High levels may increase your risk of brain and heart problems.',
        'mmse_tooltip': 'A simple test score that checks your memory and thinking. Higher scores (24-30) mean better memory function.',
        'functional_assessment_tooltip': 'How well you handle daily tasks like paying bills or taking medicine. 0 means difficulty, 10 means no problems.',
        'adl_tooltip': 'How well you do basic things like bathing, dressing, and eating. Higher scores mean more independence.',
        'memory_complaints_tooltip': 'Do you or your family notice you\'re forgetting things more often than before?',
        'behavioral_problems_tooltip': 'Unusual behaviors like aggression, wandering, or withdrawal from social activities that are new or different.',
        'confusion_tooltip': 'Difficulty thinking clearly or understanding what\'s happening around you, mixing up people or places.',
        'disorientation_tooltip': 'Getting confused about the date, time, location, or not recognizing familiar people or places.',
        'personality_changes_tooltip': 'Becoming more suspicious, fearful, angry, or withdrawn when this wasn\'t part of your usual personality.',
        'task_completion_tooltip': 'Struggling to finish familiar activities like cooking, driving to known places, or managing household duties.',
        'forgetfulness_tooltip': 'Frequently forgetting recent conversations, appointments, or events, and this is getting worse over time.',
         'mri_brain_scan_analysis': 'MRI Brain Scan Analysis',
        'mri_analysis_desc': 'Upload your brain MRI scan for comprehensive AI-powered analysis and dementia risk assessment',
        'upload_guidelines': 'Upload Guidelines',
        'supported_formats': 'Supported Formats',
        'image_size': 'Image Size',
        'max_10mb': 'Maximum 10MB',
        'image_quality': 'Image Quality',
        'high_resolution': 'High resolution preferred',
        'scan_type': 'Scan Type',
        't1_t2_flair': 'T1, T2, or FLAIR sequences',
        'patient_id': 'Patient ID',
        'auto_generated': 'Auto-generated',
        'education_years': 'Education Years',
        'socioeconomic_status': 'Socioeconomic Status',
        'cdr_score': 'CDR Score',
        'cdr': 'CDR',
        'etiv': 'eTIV (Estimated Total Intracranial Volume)',
        'nwbv': 'nWBV (Normalized Whole Brain Volume)',
        'asf': 'ASF (Atlas Scaling Factor)',
        'mri_scan_upload': 'MRI Scan Upload',
        'drag_drop_mri': 'Drag and drop your MRI scan here',
        'or_click_browse': 'or click to browse files',
        'supports_formats': 'Supports: JPEG, PNG, DICOM, NIfTI',
        'clinical_notes': 'Additional Clinical Notes',
        'clinical_notes_placeholder': 'Enter any relevant clinical observations or symptoms...',
        'analyze_mri_scan': 'Analyze MRI Scan',
        'reset_form': 'Reset Form',
        'ai_powered_analysis': 'AI-Powered Analysis',
        'ai_analysis_desc': 'Advanced machine learning algorithms for accurate dementia detection',
        'secure_confidential': 'Secure & Confidential',
        'secure_desc': 'Your medical data is encrypted and handled with utmost privacy',
        'fast_results': 'Fast Results',
        'fast_results_desc': 'Get comprehensive analysis reports within minutes',
        'medical_disclaimer': 'Medical Disclaimer',
        'analyzing': 'Analyzing',
        'invalid_file_format': 'Invalid file format. Please upload JPEG, PNG, DICOM, or NIfTI files.',
        'file_too_large': 'File size exceeds 10MB limit.',
        
        # Tooltips for MRI parameters
        'patient_id_tooltip': 'Unique identifier automatically generated based on your username and current date',
        'mri_age_tooltip': 'Patient age at the time of MRI scan. Age is crucial as brain volume naturally decreases with aging.',
        'education_years_tooltip': 'Total years of formal education completed. Higher education is associated with greater cognitive reserve.',
        'ses_tooltip': 'Socioeconomic status on a scale of 1-5, where 1 is highest and 5 is lowest. This reflects access to healthcare and lifestyle factors.',
        'etiv_tooltip': 'Estimated Total Intracranial Volume (in mm³). This measures the maximum brain size and remains constant throughout adulthood. Normal range: 1100-2000 mm³.',
        'nwbv_tooltip': 'Normalized Whole Brain Volume (ratio). This represents brain volume relative to intracranial volume. Values typically range from 0.6-0.9. Lower values may indicate atrophy.',
        'asf_tooltip': 'Atlas Scaling Factor. A normalization factor used to compare brain sizes across individuals. Normal range: 0.9-1.8.',
    },
    'ta': {
        'app_name': 'நியூரோடிடெக்ட்',
        'welcome_back': 'மீண்டும் வரவேற்கிறோம்',
        'dashboard': 'கட்டுப்பாட்டு பலகை',
        'profile': 'சுயவிவரம்',
        'cognitive_tests': 'அறிவாற்றல் சோதனைகள்',
        'mri_upload': 'எம்ஆர்ஐ பதிவேற்றம்',
        'results': 'முடிவுகள்',
        'history': 'வரலாறு',
        'logout': 'வெளியேறு',
        'notifications': 'அறிவிப்புகள்',
        'total_tests': 'மொத்த சோதனைகள்',
        'risk_level': 'இடர் நிலை',
        'avg_score': 'சராசரி மதிப்பெண்',
        'last_test': 'கடைசி சோதனை',
        'ago': 'முன்பு',
        'quick_actions': 'விரைவு செயல்கள்',
        'start_cognitive_test': 'அறிவாற்றல் சோதனையைத் தொடங்கு',
        'cognitive_test_desc': 'உங்கள் விரிவான அறிவாற்றல் மதிப்பீட்டைத் தொடங்குங்கள்',
        'upload_mri_scan': 'எம்ஆர்ஐ ஸ்கேனை பதிவேற்றவும்',
        'mri_scan_desc': 'AI-இயங்கும் பகுப்பாய்வுக்கு மூளை ஸ்கேனை பதிவேற்றவும்',
        'view_results': 'முடிவுகளைக் காண்க',
        'view_results_desc': 'உங்கள் சமீபத்திய மதிப்பீட்டு முடிவுகளைச் சரிபார்க்கவும்',
        'test_history': 'சோதனை வரலாறு',
        'test_history_desc': 'கடந்த கால மதிப்பீடுகள் மற்றும் போக்குகளைப் பார்க்கவும்',
        'recent_tests': 'சமீபத்திய சோதனைகள்',
        'view_all': 'அனைத்தையும் காண்க',
        'mark_all_read': 'அனைத்தையும் படித்ததாகக் குறிக்கவும்',
        'daily_health_tip': '💡 தினசரி ஆரோக்கிய குறிப்பு',
        'health_tip_text': 'வழக்கமான மனப் பயிற்சிகள் மற்றும் ஒமேகா-3 கொழுப்பு அமிலங்கள் நிறைந்த சீரான உணவு அறிவாற்றல் செயல்பாட்டை பராமரிக்க உதவும். இன்று புதிர் விளையாட்டுகள், தியானம் அல்லது ஒரு புதிய திறமையைக் கற்றுக்கொள்ளுங்கள்!',
        'low_risk': 'குறைந்த ஆபத்து',
        'high_risk': 'அதிக ஆபத்து',
        'score': 'மதிப்பெண்',
        'view_details': 'விவரங்களைக் காண்க',
        'patient_name': 'நோயாளியின் முழு பெயர்',
        'age': 'வயது',
        'gender': 'பாலினம்',
        'male': 'ஆண்',
        'female': 'பெண்',
        'select_gender': 'பாலினத்தைத் தேர்ந்தெடுக்கவும்',
        'submit': 'சமர்ப்பிக்கவும்',
        'generate_assessment': 'இடர் மதிப்பீட்டை உருவாக்கவும்',
        'comprehensive_assessment': 'விரிவான இடர் மதிப்பீடு',
        'patient_info': 'நோயாளி தகவல்',
        'medical_history': 'மருத்துவ மற்றும் குடும்ப வரலாறு',
        'lifestyle': 'வாழ்க்கை முறை மற்றும் உடல் ஆரோக்கியம்',
        'download_pdf': 'PDF அறிக்கையைப் பதிவிறக்கவும்',
        'download_docx': 'Word அறிக்கையைப் பதிவிறக்கவும்',
        'prediction_result': 'கணிப்பு முடிவு',
        'no_data': 'தரவு கிடைக்கவில்லை',
        'update_profile': 'சுயவிவரத்தைப் புதுப்பிக்கவும்',
        'email': 'மின்னஞ்சல்',
        'save_changes': 'மாற்றங்களைச் சேமிக்கவும்',
        'language': 'மொழி',
        'select_language': 'மொழியைத் தேர்ந்தெடுக்கவும்',
        'login': 'உள்நுழைய',
        'signup': 'பதிவு செய்ய',
        'username': 'பயனர்பெயர்',
        'password': 'கடவுச்சொல்',
        'start_assessment': 'மதிப்பீட்டைத் தொடங்கவும்',
        'monitor_health': 'எங்கள் AI-இயங்கும் நுண்ணறிவுகளுடன் உங்கள் அறிவாற்றல் ஆரோக்கிய பயணத்தை கண்காணிக்கவும்',
        'back_to_dashboard': 'கட்டுப்பாட்டு பலகைக்குத் திரும்பவும்',
        'complete_form_desc': 'AI-இயங்கும் அல்சைமர் இடர் கணிப்பை உருவாக்க பின்வரும் படிவத்தை நிரப்பவும்',
        'patient_info_desc': 'அடிப்படை மக்கள்தொகை மற்றும் அடையாள விவரங்கள்',
        'medical_history_desc': 'முந்தைய நிலைமைகள் மற்றும் குடும்ப மருத்துவ பின்னணி',
        'lifestyle_desc': 'தினசரி பழக்கவழக்கங்கள் மற்றும் உடல் ஆரோக்கிய குறிகாட்டிகள்',
        'vital_signs': 'முக்கிய அறிகுறிகள் & ஆய்வக முடிவுகள்',
        'vital_signs_desc': 'இரத்த அழுத்தம், கொழுப்பு மற்றும் வளர்சிதை மாற்ற குறிப்பான்கள்',
        'cognitive_assessment': 'அறிவாற்றல் & செயல்பாட்டு மதிப்பீடு',
        'cognitive_assessment_desc': 'மன திறன் மற்றும் தினசரி செயல்பாட்டு மதிப்பெண்கள்',
        'clinical_symptoms': 'மருத்துவ அறிகுறிகள் & நடத்தை குறிகாட்டிகள்',
        'clinical_symptoms_desc': 'கவனிக்கப்பட்ட அறிகுறிகள் மற்றும் நடத்தை மாற்றங்கள்',
        'enter_patient_name': 'நோயாளியின் முழு பெயரை உள்ளிடவும்',
        'years': 'ஆண்டுகள்',
        'ethnicity': 'இனம்',
        'select_ethnicity': 'இனத்தைத் தேர்ந்தெடுக்கவும்',
        'other': 'மற்றவை',
        'asian': 'ஆசிய',
        'black': 'கறுப்பு',
        'white': 'வெள்ளை',
        'education_level': 'கல்வி நிலை',
        'select_level': 'நிலையைத் தேர்ந்தெடுக்கவும்',
        'no_schooling': 'பள்ளிக் கல்வி இல்லை',
        'primary_education': 'முதன்மைக் கல்வி',
        'secondary_education': 'இடைநிலைக் கல்வி',
        'higher_education': 'உயர்கல்வி',
        'family_history_alz': 'அல்சைமரின் குடும்ப வரலாறு',
        'family_history': 'குடும்ப வரலாறு',
        'yes': 'ஆம்',
        'no': 'இல்லை',
        'diabetes': 'நீரிழிவு',
        'hypertension': 'உயர் இரத்த அழுத்தம்',
        'cardiovascular_disease': 'இதய நோய்',
        'depression': 'மனச்சோர்வு',
        'head_injury': 'தலை காயம்',
        'bmi': 'பிஎம்ஐ',
        'physical_activity': 'உடல் செயல்பாட்டு நிலை',
        'smoking_status': 'புகைபிடித்தல் நிலை',
        'smoking': 'புகைபிடித்தல்',
        'non_smoker': 'புகைபிடிக்காதவர்',
        'smoker': 'புகைபிடிப்பவர்',
        'alcohol_consumption': 'மது அருந்துதல்',
        'diet_quality': 'உணவு தர மதிப்பெண்',
        'sleep_quality': 'தூக்க தர மதிப்பெண்',
        'systolic_bp': 'சிஸ்டாலிக் இரத்த அழுத்தம்',
        'diastolic_bp': 'டயஸ்டாலிக் இரத்த அழுத்தம்',
        'total_cholesterol': 'மொத்த கொழுப்பு',
        'ldl_cholesterol': 'எல்டிஎல் கொழுப்பு',
        'hdl_cholesterol': 'எச்டிஎல் கொழுப்பு',
        'triglycerides': 'ட்ரைகிளிசரைடுகள்',
        'mmse_score': 'எம்எம்எஸ்இ மதிப்பெண்',
        'mmse': 'எம்எம்எஸ்இ',
        'functional_assessment': 'செயல்பாட்டு மதிப்பீட்டு மதிப்பெண்',
        'adl_score': 'ஏடிஎல் மதிப்பெண்',
        'adl': 'ஏடிஎல்',
        'memory_complaints': 'நினைவாற்றல் புகார்கள்',
        'behavioral_problems': 'நடத்தை பிரச்சனைகள்',
        'confusion': 'குழப்பம்',
        'disorientation': 'திசைதிருப்பல்',
        'personality_changes': 'ஆளுமை மாற்றங்கள்',
        'difficulty_completing_tasks': 'பணிகளை முடிப்பதில் சிரமம்',
        'forgetfulness': 'மறதி',
        'ai_powered_platform': 'AI-இயங்கும் முன்கூட்டியே கண்டறியும் தளம்',
        'disclaimer_text': 'இந்த கருவி ஸ்கிரீனிங் நோக்கங்களுக்காக மட்டுமே மற்றும் தொழில்முறை மருத்துவ நோயறிதலை மாற்றாது.',
        'age_tooltip': 'உங்கள் தற்போதைய வயது முக்கியம், ஏனெனில் அல்சைமர் ஆபத்து வயதாகும்போது அதிகரிக்கிறது, குறிப்பாக 65 க்குப் பிறகு.',
        'gender_tooltip': 'பெண்களுக்கு அல்சைமர் ஆபத்து சற்று அதிகம், ஏனெனில் அவர்கள் பொதுவாக ஆண்களை விட நீண்ட காலம் வாழ்கிறார்கள்.',
        'ethnicity_tooltip': 'சில இன குழுக்களுக்கு மரபணு மற்றும் சுகாதார காரணிகளால் வெவ்வேறு ஆபத்து நிலைகள் உள்ளன.',
        'education_tooltip': 'உயர் கல்வி மூளை வலிமையை உருவாக்க உதவுகிறது, இது வாழ்க்கையின் பிற்பகுதியில் நினைவக பிரச்சனைகளிலிருந்து பாதுகாக்கலாம்.',
        'family_history_tooltip': 'உங்கள் பெற்றோர் அல்லது உடன்பிறப்புகளுக்கு அல்சைமர் இருந்தால், அதை உருவாக்கும் வாய்ப்பு உங்களுக்கு அதிகம்.',
        'diabetes_tooltip': 'நீரிழிவு நோயிலிருந்து அதிக இரத்த சர்க்கரை காலப்போக்கில் உங்கள் மூளையை சேதப்படுத்தலாம் மற்றும் நினைவக பிரச்சனை அபாயங்களை அதிகரிக்கலாம்.',
        'hypertension_tooltip': 'உயர் இரத்த அழுத்தம் உங்கள் மூளையில் உள்ள இரத்த நாளங்களை சேதப்படுத்தலாம், இரத்த ஓட்டத்தை குறைக்கலாம் மற்றும் நினைவகத்தை பாதிக்கலாம்.',
        'cardiovascular_tooltip': 'இதய பிரச்சனைகள் உங்கள் மூளைக்கு இரத்த ஓட்டத்தை பாதிக்கலாம், இது நினைவக பிரச்சனைகளின் ஆபத்தை அதிகரிக்கலாம்.',
        'depression_tooltip': 'நீண்டகால மனச்சோர்வு உங்கள் மூளை ஆரோக்கியத்தை பாதிக்கலாம் மற்றும் அதிக அல்சைமர் ஆபத்துடன் இணைக்கப்படலாம்.',
        'head_injury_tooltip': 'கடுமையான தலை காயங்கள் அல்லது மீண்டும் மீண்டும் மூளையதிர்ச்சி பின்னர் நினைவக பிரச்சனைகளை உருவாக்கும் ஆபத்தை அதிகரிக்கலாம்.',
        'bmi_tooltip': 'உயரத்துடன் ஒப்பிடும்போது உங்கள் எடை. மிகவும் அதிக எடை அல்லது குறைவான எடை மூளை ஆரோக்கியத்தை பாதிக்கலாம்.',
        'physical_activity_tooltip': 'நீங்கள் எவ்வளவு உடற்பயிற்சி செய்கிறீர்கள். வழக்கமான செயல்பாடு உங்கள் மூளையை ஆரோக்கியமாக வைக்கிறது. 0 என்றால் உடற்பயிற்சி இல்லை, 10 என்றால் தினசரி பயிற்சிகள்.',
        'smoking_tooltip': 'புகைபிடித்தல் உங்கள் மூளையில் உள்ள இரத்த நாளங்களை சேதப்படுத்துகிறது மற்றும் நினைவக பிரச்சனைகளின் ஆபத்தை கணிசமாக அதிகரிக்கிறது.',
        'alcohol_tooltip': 'நீங்கள் வாரத்திற்கு எத்தனை மதுபான பானங்கள் அருந்துகிறீர்கள். அதிக அளவு குடிப்பது காலப்போக்கில் உங்கள் மூளையை சேதப்படுத்தலாம்.',
        'diet_quality_tooltip': 'நீங்கள் எவ்வளவு ஆரோக்கியமாக சாப்பிடுகிறீர்கள். அதிக பழங்கள், காய்கறிகள் மற்றும் மீன் மூளை ஆரோக்கியத்திற்கு உதவுகின்றன. 0 என்பது மோசமான உணவு, 10 என்பது சிறந்தது.',
        'sleep_quality_tooltip': 'நீங்கள் எவ்வளவு நன்றாக தூங்குகிறீர்கள். நல்ல தூக்கம் உங்கள் மூளையிலிருந்து நச்சுகளை சுத்தம் செய்ய உதவுகிறது. 0 மிகவும் மோசமான தூக்கம், 10 சிறந்தது.',
        'systolic_bp_tooltip': 'உங்கள் இரத்த அழுத்த வாசிப்பில் மேல் எண். அதிக எண்கள் உங்கள் மூளையில் உள்ள இரத்த நாளங்களை சேதப்படுத்தலாம்.',
        'diastolic_bp_tooltip': 'உங்கள் இரத்த அழுத்த வாசிப்பில் கீழ் எண். இதயம் துடிப்புகளுக்கு இடையில் ஓய்வெடுக்கும்போது அழுத்தத்தை காட்டுகிறது.',
        'cholesterol_total_tooltip': 'உங்கள் இரத்தத்தில் உள்ள அனைத்து கொழுப்பு. மிக அதிக அளவுகள் உங்கள் மூளைக்கு உணவளிக்கும் இரத்த நாளங்களை சேதப்படுத்தலாம்.',
        'ldl_tooltip': 'இரத்த நாளங்களை அடைக்கக்கூடிய கெட்ட கொழுப்பு. குறைந்த எண்கள் மூளை ஆரோக்கியத்திற்கு சிறந்தவை.',
        'hdl_tooltip': 'கெட்ட கொழுப்பை அகற்ற உதவும் நல்ல கொழுப்பு. அதிக எண்கள் உங்கள் மூளையை பாதுகாக்கின்றன.',
        'triglycerides_tooltip': 'உங்கள் இரத்தத்தில் உள்ள ஒரு வகை கொழுப்பு. அதிக அளவுகள் மூளை மற்றும் இதய பிரச்சனைகளின் ஆபத்தை அதிகரிக்கலாம்.',
        'mmse_tooltip': 'உங்கள் நினைவகம் மற்றும் சிந்தனையை சரிபார்க்கும் எளிய சோதனை மதிப்பெண். அதிக மதிப்பெண்கள் (24-30) சிறந்த நினைவக செயல்பாட்டை குறிக்கின்றன.',
        'functional_assessment_tooltip': 'பில்களை செலுத்துதல் அல்லது மருந்து எடுத்தல் போன்ற தினசரி பணிகளை நீங்கள் எவ்வளவு நன்றாக கையாளுகிறீர்கள். 0 என்றால் சிரமம், 10 என்றால் பிரச்சனை இல்லை.',
        'adl_tooltip': 'குளித்தல், ஆடை அணிதல் மற்றும் உண்ணுதல் போன்ற அடிப்படை விஷயங்களை நீங்கள் எவ்வளவு நன்றாக செய்கிறீர்கள். அதிக மதிப்பெண்கள் அதிக சுதந்திரத்தை குறிக்கின்றன.',
        'memory_complaints_tooltip': 'நீங்கள் அல்லது உங்கள் குடும்பம் முன்பை விட அடிக்கடி விஷயங்களை மறக்கிறீர்கள் என்பதை கவனிக்கிறீர்களா?',
        'behavioral_problems_tooltip': 'ஆக்கிரமிப்பு, அலைதல் அல்லது சமூக நடவடிக்கைகளிலிருந்து விலகல் போன்ற புதிய அல்லது வித்தியாசமான அசாதாரண நடத்தைகள்.',
        'confusion_tooltip': 'தெளிவாக சிந்திக்க அல்லது உங்களைச் சுற்றி என்ன நடக்கிறது என்பதைப் புரிந்துகொள்வதில் சிரமம், மக்கள் அல்லது இடங்களை குழப்புதல்.',
        'disorientation_tooltip': 'தேதி, நேரம், இடம் பற்றி குழப்பமடைதல் அல்லது பழக்கமான நபர்கள் அல்லது இடங்களை அடையாளம் காணாமல் இருத்தல்.',
        'personality_changes_tooltip': 'இது உங்கள் வழக்கமான ஆளுமையின் ஒரு பகுதியாக இல்லாதபோது அதிக சந்தேகம், பயம், கோபம் அல்லது விலகல்.',
        'task_completion_tooltip': 'சமையல், அறிந்த இடங்களுக்கு வாகனம் ஓட்டுதல் அல்லது வீட்டு கடமைகளை நிர்வகித்தல் போன்ற பழக்கமான செயல்பாடுகளை முடிக்க போராடுதல்.',
        'forgetfulness_tooltip': 'சமீபத்திய உரையாடல்கள், சந்திப்புகள் அல்லது நிகழ்வுகளை அடிக்கடி மறப்பது, இது காலப்போக்கில் மோசமாகி வருகிறது.',
         'mri_brain_scan_analysis': 'எம்ஆர்ஐ மூளை ஸ்கேன் பகுப்பாய்வு',
        'mri_analysis_desc': 'விரிவான AI-இயங்கும் பகுப்பாய்வு மற்றும் டிமென்ஷியா இடர் மதிப்பீட்டிற்காக உங்கள் மூளை எம்ஆர்ஐ ஸ்கேனை பதிவேற்றவும்',
        'upload_guidelines': 'பதிவேற்ற வழிகாட்டுதல்கள்',
        'supported_formats': 'ஆதரிக்கப்படும் வடிவங்கள்',
        'image_size': 'படத்தின் அளவு',
        'max_10mb': 'அதிகபட்சம் 10MB',
        'image_quality': 'படத்தின் தரம்',
        'high_resolution': 'உயர் தெளிவுத்திறன் விரும்பப்படும்',
        'scan_type': 'ஸ்கேன் வகை',
        't1_t2_flair': 'T1, T2, அல்லது FLAIR வரிசைகள்',
        'patient_id': 'நோயாளி அடையாள எண்',
        'auto_generated': 'தானாக உருவாக்கப்பட்டது',
        'education_years': 'கல்வி ஆண்டுகள்',
        'socioeconomic_status': 'சமூக பொருளாதார நிலை',
        'cdr_score': 'CDR மதிப்பெண்',
        'cdr': 'CDR',
        'etiv': 'eTIV (மதிப்பிடப்பட்ட மொத்த மண்டை உள் கொள்ளளவு)',
        'nwbv': 'nWBV (இயல்பாக்கப்பட்ட முழு மூளை கொள்ளளவு)',
        'asf': 'ASF (அட்லஸ் அளவிடும் காரணி)',
        'mri_scan_upload': 'எம்ஆர்ஐ ஸ்கேன் பதிவேற்றம்',
        'drag_drop_mri': 'உங்கள் எம்ஆர்ஐ ஸ்கேனை இங்கே இழுத்து விடவும்',
        'or_click_browse': 'அல்லது கோப்புகளை உலாவ கிளிக் செய்யவும்',
        'supports_formats': 'ஆதரிக்கிறது: JPEG, PNG, DICOM, NIfTI',
        'clinical_notes': 'கூடுதல் மருத்துவ குறிப்புகள்',
        'clinical_notes_placeholder': 'ஏதேனும் தொடர்புடைய மருத்துவ அவதானிப்புகள் அல்லது அறிகுறிகளை உள்ளிடவும்...',
        'analyze_mri_scan': 'எம்ஆர்ஐ ஸ்கேனை பகுப்பாய்வு செய்யவும்',
        'reset_form': 'படிவத்தை மீட்டமைக்கவும்',
        'ai_powered_analysis': 'AI-இயங்கும் பகுப்பாய்வு',
        'ai_analysis_desc': 'துல்லியமான டிமென்ஷியா கண்டறிதலுக்கான மேம்பட்ட இயந்திர கற்றல் வழிமுறைகள்',
        'secure_confidential': 'பாதுகாப்பான & ரகசியமான',
        'secure_desc': 'உங்கள் மருத்துவ தரவு என்க்ரிப்ட் செய்யப்பட்டு மிக உயர்ந்த தனியுரிமையுடன் கையாளப்படுகிறது',
        'fast_results': 'விரைவான முடிவுகள்',
        'fast_results_desc': 'நிமிடங்களில் விரிவான பகுப்பாய்வு அறிக்கைகளைப் பெறுங்கள்',
        'medical_disclaimer': 'மருத்துவ மறுப்பு',
        'analyzing': 'பகுப்பாய்வு செய்கிறது',
        'invalid_file_format': 'தவறான கோப்பு வடிவம். JPEG, PNG, DICOM, அல்லது NIfTI கோப்புகளை பதிவேற்றவும்.',
        'file_too_large': 'கோப்பு அளவு 10MB வரம்பை மீறுகிறது.',
        
        # Tooltips
        'patient_id_tooltip': 'உங்கள் பயனர்பெயர் மற்றும் தற்போதைய தேதியின் அடிப்படையில் தானாக உருவாக்கப்படும் தனித்துவமான அடையாளங்காட்டி',
        'mri_age_tooltip': 'எம்ஆர்ஐ ஸ்கேன் நேரத்தில் நோயாளியின் வயது. வயது முக்கியமானது, ஏனெனில் மூளையின் அளவு இயற்கையாகவே வயதாகும்போது குறைகிறது.',
        'education_years_tooltip': 'முடிக்கப்பட்ட முறையான கல்வியின் மொத்த ஆண்டுகள். உயர் கல்வி அதிக அறிவாற்றல் இருப்புடன் தொடர்புடையது.',
        'ses_tooltip': '1-5 அளவில் சமூக பொருளாதார நிலை, இதில் 1 மிக உயர்ந்தது மற்றும் 5 மிகக் குறைந்தது. இது சுகாதார பராமரிப்பு மற்றும் வாழ்க்கை முறை காரணிகளுக்கான அணுகலை பிரதிபலிக்கிறது.',
        'etiv_tooltip': 'மதிப்பிடப்பட்ட மொத்த மண்டை உள் கொள்ளளவு (mm³ இல்). இது அதிகபட்ச மூளை அளவை அளவிடுகிறது மற்றும் வயது முழுவதும் மாறாமல் இருக்கும். சாதாரண வரம்பு: 1100-2000 mm³.',
        'nwbv_tooltip': 'இயல்பாக்கப்பட்ட முழு மூளை கொள்ளளவு (விகிதம்). இது மண்டை உள் கொள்ளளவுடன் ஒப்பிடும்போது மூளை கொள்ளளவை குறிக்கிறது. மதிப்புகள் பொதுவாக 0.6-0.9 வரம்பில் இருக்கும். குறைந்த மதிப்புகள் சிதைவைக் குறிக்கலாம்.',
        'asf_tooltip': 'அட்லஸ் அளவிடும் காரணி. தனிநபர்களின் மூளை அளவுகளை ஒப்பிட பயன்படுத்தப்படும் இயல்பாக்க காரணி. சாதாரண வரம்பு: 0.9-1.8.',
    },
    'hi': {
        'app_name': 'न्यूरोडिटेक्ट',
        'welcome_back': 'वापसी पर स्वागत है',
        'dashboard': 'डैशबोर्ड',
        'profile': 'प्रोफ़ाइल',
        'cognitive_tests': 'संज्ञानात्मक परीक्षण',
        'mri_upload': 'एमआरआई अपलोड',
        'results': 'परिणाम',
        'history': 'इतिहास',
        'logout': 'लॉग आउट',
        'notifications': 'सूचनाएं',
        'total_tests': 'कुल परीक्षण',
        'risk_level': 'जोखिम स्तर',
        'avg_score': 'औसत स्कोर',
        'last_test': 'अंतिम परीक्षण',
        'ago': 'पहले',
        'quick_actions': 'त्वरित कार्य',
        'start_cognitive_test': 'संज्ञानात्मक परीक्षण शुरू करें',
        'cognitive_test_desc': 'अपना व्यापक संज्ञानात्मक मूल्यांकन शुरू करें',
        'upload_mri_scan': 'एमआरआई स्कैन अपलोड करें',
        'mri_scan_desc': 'एआई-संचालित विश्लेषण के लिए ब्रेन स्कैन अपलोड करें',
        'view_results': 'परिणाम देखें',
        'view_results_desc': 'अपने नवीनतम मूल्यांकन परिणाम जांचें',
        'test_history': 'परीक्षण इतिहास',
        'test_history_desc': 'सभी पिछले मूल्यांकन और रुझान देखें',
        'recent_tests': 'हाल के परीक्षण',
        'view_all': 'सभी देखें',
        'mark_all_read': 'सभी को पढ़ा हुआ चिह्नित करें',
        'daily_health_tip': '💡 दैनिक स्वास्थ्य टिप',
        'health_tip_text': 'नियमित मानसिक अभ्यास और ओमेगा-3 फैटी एसिड से भरपूर संतुलित आहार संज्ञानात्मक कार्य को बनाए रखने में मदद कर सकता है। आज पहेली खेल, ध्यान या एक नया कौशल सीखने का प्रयास करें!',
        'low_risk': 'कम जोखिम',
        'high_risk': 'उच्च जोखिम',
        'score': 'स्कोर',
        'view_details': 'विवरण देखें',
        'patient_name': 'रोगी का पूरा नाम',
        'age': 'आयु',
        'gender': 'लिंग',
        'male': 'पुरुष',
        'female': 'महिला',
        'select_gender': 'लिंग चुनें',
        'submit': 'जमा करें',
        'generate_assessment': 'जोखिम मूल्यांकन उत्पन्न करें',
        'comprehensive_assessment': 'व्यापक जोखिम मूल्यांकन',
        'patient_info': 'रोगी की जानकारी',
        'medical_history': 'चिकित्सा और पारिवारिक इतिहास',
        'lifestyle': 'जीवनशैली और शारीरिक स्वास्थ्य',
        'download_pdf': 'पीडीएफ रिपोर्ट डाउनलोड करें',
        'download_docx': 'वर्ड रिपोर्ट डाउनलोड करें',
        'prediction_result': 'भविष्यवाणी परिणाम',
        'no_data': 'कोई डेटा उपलब्ध नहीं',
        'update_profile': 'प्रोफ़ाइल अपडेट करें',
        'email': 'ईमेल',
        'save_changes': 'परिवर्तन सहेजें',
        'language': 'भाषा',
        'select_language': 'भाषा चुनें',
        'login': 'लॉगिन',
        'signup': 'साइन अप',
        'username': 'उपयोगकर्ता नाम',
        'password': 'पासवर्ड',
        'start_assessment': 'मूल्यांकन शुरू करें',
        'monitor_health': 'हमारे एआई-संचालित अंतर्दृष्टि के साथ अपनी संज्ञानात्मक स्वास्थ्य यात्रा की निगरानी करें',
        'back_to_dashboard': 'डैशबोर्ड पर वापस जाएं',
        'complete_form_desc': 'एआई-संचालित अल्जाइमर जोखिम भविष्यवाणी उत्पन्न करने के लिए निम्नलिखित फॉर्म भरें',
        'patient_info_desc': 'बुनियादी जनसांख्यिकीय और पहचान विवरण',
        'medical_history_desc': 'पिछली स्थितियाँ और पारिवारिक चिकित्सा पृष्ठभूमि',
        'lifestyle_desc': 'दैनिक आदतें और शारीरिक स्वास्थ्य संकेतक',
        'vital_signs': 'महत्वपूर्ण संकेत और प्रयोगशाला परिणाम',
        'vital_signs_desc': 'रक्तचाप, कोलेस्ट्रॉल और चयापचय मार्कर',
        'cognitive_assessment': 'संज्ञानात्मक और कार्यात्मक मूल्यांकन',
        'cognitive_assessment_desc': 'मानसिक क्षमता और दैनिक कामकाज स्कोर',
        'clinical_symptoms': 'नैदानिक लक्षण और व्यवहार संकेतक',
        'clinical_symptoms_desc': 'देखे गए लक्षण और व्यवहार परिवर्तन',
        'enter_patient_name': 'रोगी का पूरा नाम दर्ज करें',
        'years': 'वर्ष',
        'ethnicity': 'जातीयता',
        'select_ethnicity': 'जातीयता चुनें',
        'other': 'अन्य',
        'asian': 'एशियाई',
        'black': 'काला',
        'white': 'सफेद',
        'education_level': 'शिक्षा स्तर',
        'select_level': 'स्तर चुनें',
        'no_schooling': 'कोई स्कूली शिक्षा नहीं',
        'primary_education': 'प्राथमिक शिक्षा',
        'secondary_education': 'माध्यमिक शिक्षा',
        'higher_education': 'उच्च शिक्षा',
        'family_history_alz': 'अल्जाइमर का पारिवारिक इतिहास',
        'family_history': 'पारिवारिक इतिहास',
        'yes': 'हाँ',
        'no': 'नहीं',
        'diabetes': 'मधुमेह',
        'hypertension': 'उच्च रक्तचाप',
        'cardiovascular_disease': 'हृदय रोग',
        'depression': 'अवसाद',
        'head_injury': 'सिर की चोट',
        'bmi': 'बीएमआई',
        'physical_activity': 'शारीरिक गतिविधि स्तर',
        'smoking_status': 'धूम्रपान की स्थिति',
        'smoking': 'धूम्रपान',
        'non_smoker': 'धूम्रपान न करने वाला',
        'smoker': 'धूम्रपान करने वाला',
        'alcohol_consumption': 'शराब की खपत',
        'diet_quality': 'आहार गुणवत्ता स्कोर',
        'sleep_quality': 'नींद की गुणवत्ता स्कोर',
        'systolic_bp': 'सिस्टोलिक रक्तचाप',
        'diastolic_bp': 'डायस्टोलिक रक्तचाप',
        'total_cholesterol': 'कुल कोलेस्ट्रॉल',
        'ldl_cholesterol': 'एलडीएल कोलेस्ट्रॉल',
        'hdl_cholesterol': 'एचडीएल कोलेस्ट्रॉल',
        'triglycerides': 'ट्राइग्लिसराइड्स',
        'mmse_score': 'एमएमएसई स्कोर',
        'mmse': 'एमएमएसई',
        'functional_assessment': 'कार्यात्मक मूल्यांकन स्कोर',
        'adl_score': 'एडीएल स्कोर',
        'adl': 'एडीएल',
        'memory_complaints': 'स्मृति शिकायतें',
        'behavioral_problems': 'व्यवहार संबंधी समस्याएं',
        'confusion': 'भ्रम',
        'disorientation': 'दिशाहीनता',
        'personality_changes': 'व्यक्तित्व परिवर्तन',
        'difficulty_completing_tasks': 'कार्य पूरा करने में कठिनाई',
        'forgetfulness': 'भूलने की बीमारी',
        'ai_powered_platform': 'एआई-संचालित प्रारंभिक पहचान मंच',
        'disclaimer_text': 'यह उपकरण केवल स्क्रीनिंग उद्देश्यों के लिए है और पेशेवर चिकित्सा निदान की जगह नहीं लेता है।',
        'age_tooltip': 'आपकी वर्तमान आयु मायने रखती है क्योंकि अल्जाइमर का जोखिम बढ़ती उम्र के साथ बढ़ता है, खासकर 65 के बाद।',
        'gender_tooltip': 'महिलाओं में अल्जाइमर का जोखिम थोड़ा अधिक होता है, आंशिक रूप से क्योंकि वे आम तौर पर पुरुषों की तुलना में अधिक समय तक जीवित रहती हैं।',
        'ethnicity_tooltip': 'आनुवंशिक और स्वास्थ्य कारकों के कारण कुछ जातीय समूहों में अल्जाइमर के लिए विभिन्न जोखिम स्तर होते हैं।',
        'education_tooltip': 'उच्च शिक्षा मस्तिष्क की शक्ति बनाने में मदद करती है, जो जीवन में बाद में स्मृति समस्याओं से बचा सकती है।',
        'family_history_tooltip': 'यदि आपके माता-पिता या भाई-बहनों को अल्जाइमर था, तो आपको इसके विकसित होने की अधिक संभावना हो सकती है।',
        'diabetes_tooltip': 'मधुमेह से उच्च रक्त शर्करा समय के साथ आपके मस्तिष्क को नुकसान पहुंचा सकती है और स्मृति समस्या जोखिम बढ़ा सकती है।',
        'hypertension_tooltip': 'उच्च रक्तचाप आपके मस्तिष्क में रक्त वाहिकाओं को नुकसान पहुंचा सकता है, रक्त प्रवाह को कम कर सकता है और स्मृति को प्रभावित कर सकता है।',
        'cardiovascular_tooltip': 'हृदय की समस्याएं आपके मस्तिष्क में रक्त प्रवाह को प्रभावित कर सकती हैं, जो स्मृति मुद्दों के जोखिम को बढ़ा सकती हैं।',
        'depression_tooltip': 'दीर्घकालिक अवसाद आपके मस्तिष्क स्वास्थ्य को प्रभावित कर सकता है और उच्च अल्जाइमर जोखिम से जुड़ा हो सकता है।',
        'head_injury_tooltip': 'गंभीर सिर की चोटें या बार-बार की मस्तिष्क हिलना बाद में स्मृति समस्याओं के विकास के आपके जोखिम को बढ़ा सकता है।',
        'bmi_tooltip': 'ऊंचाई की तुलना में आपका वजन। बहुत अधिक वजन या कम वजन होना मस्तिष्क स्वास्थ्य को प्रभावित कर सकता है।',
        'physical_activity_tooltip': 'आप कितना व्यायाम करते हैं। नियमित गतिविधि आपके मस्तिष्क को स्वस्थ रखती है। 0 का मतलब कोई व्यायाम नहीं, 10 का मतलब दैनिक कसरत।',
        'smoking_tooltip': 'धूम्रपान आपके मस्तिष्क में रक्त वाहिकाओं को नुकसान पहुंचाता है और स्मृति समस्याओं के आपके जोखिम को काफी बढ़ाता है।',
        'alcohol_tooltip': 'आप प्रति सप्ताह कितने मादक पेय लेते हैं। भारी शराब पीने से समय के साथ आपके मस्तिष्क को नुकसान हो सकता है।',
        'diet_quality_tooltip': 'आप कितना स्वस्थ खाते हैं। अधिक फल, सब्जियां और मछली मस्तिष्क स्वास्थ्य में मदद करते हैं। 0 खराब आहार है, 10 उत्कृष्ट है।',
        'sleep_quality_tooltip': 'आप कितनी अच्छी नींद लेते हैं। अच्छी नींद आपके मस्तिष्क से विषाक्त पदार्थों को साफ करने में मदद करती है। 0 बहुत खराब नींद है, 10 उत्कृष्ट है।',
        'systolic_bp_tooltip': 'आपके रक्तचाप पढ़ने में शीर्ष संख्या। उच्च संख्या आपके मस्तिष्क में रक्त वाहिकाओं को नुकसान पहुंचा सकती है।',
        'diastolic_bp_tooltip': 'आपके रक्तचाप पढ़ने में नीचे की संख्या। यह दिखाता है कि जब आपका दिल धड़कनों के बीच आराम करता है तो दबाव।',
        'cholesterol_total_tooltip': 'आपके रक्त में सभी कोलेस्ट्रॉल। बहुत उच्च स्तर आपके मस्तिष्क को खिलाने वाली रक्त वाहिकाओं को नुकसान पहुंचा सकते हैं।',
        'ldl_tooltip': '\'खराब\' कोलेस्ट्रॉल जो रक्त वाहिकाओं को अवरुद्ध कर सकता है। मस्तिष्क स्वास्थ्य के लिए कम संख्या बेहतर है।',
        'hdl_tooltip': '\'अच्छा\' कोलेस्ट्रॉल जो खराब कोलेस्ट्रॉल को हटाने में मदद करता है। उच्च संख्या आपके मस्तिष्क की रक्षा करती है।',
        'triglycerides_tooltip': 'आपके रक्त में एक प्रकार की वसा। उच्च स्तर मस्तिष्क और हृदय समस्याओं के आपके जोखिम को बढ़ा सकते हैं।',
        'mmse_tooltip': 'एक सरल परीक्षण स्कोर जो आपकी स्मृति और सोच की जांच करता है। उच्च स्कोर (24-30) का मतलब बेहतर स्मृति कार्य है।',
        'functional_assessment_tooltip': 'आप बिल का भुगतान करने या दवा लेने जैसे दैनिक कार्यों को कितनी अच्छी तरह संभालते हैं। 0 का मतलब कठिनाई है, 10 का मतलब कोई समस्या नहीं है।',
        'adl_tooltip': 'आप नहाने, कपड़े पहनने और खाने जैसी बुनियादी चीजें कितनी अच्छी तरह करते हैं। उच्च स्कोर का मतलब अधिक स्वतंत्रता है।',
        'memory_complaints_tooltip': 'क्या आप या आपका परिवार नोटिस करता है कि आप पहले की तुलना में अधिक बार चीजें भूल रहे हैं?',
        'behavioral_problems_tooltip': 'आक्रामकता, भटकना, या सामाजिक गतिविधियों से हटना जैसे असामान्य व्यवहार जो नए या अलग हैं।',
        'confusion_tooltip': 'स्पष्ट रूप से सोचने में कठिनाई या यह समझने में कि आपके आसपास क्या हो रहा है, लोगों या स्थानों को मिलाना।',
        'disorientation_tooltip': 'तारीख, समय, स्थान के बारे में भ्रमित होना, या परिचित लोगों या स्थानों को पहचान नहीं पाना।',
        'personality_changes_tooltip': 'अधिक संदिग्ध, भयभीत, क्रोधित, या वापस लेना जब यह आपके सामान्य व्यक्तित्व का हिस्सा नहीं था।',
        'task_completion_tooltip': 'परिचित गतिविधियों को समाप्त करने के लिए संघर्ष करना जैसे खाना पकाना, ज्ञात स्थानों पर ड्राइविंग, या घरेलू कर्तव्यों का प्रबंधन।',
        'forgetfulness_tooltip': 'हाल की बातचीत, नियुक्तियों, या घटनाओं को अक्सर भूलना, और यह समय के साथ खराब हो रहा है।',
          'mri_brain_scan_analysis': 'एमआरआई मस्तिष्क स्कैन विश्लेषण',
        'mri_analysis_desc': 'व्यापक एआई-संचालित विश्लेषण और मनोभ्रंश जोखिम मूल्यांकन के लिए अपना मस्तिष्क एमआरआई स्कैन अपलोड करें',
        'upload_guidelines': 'अपलोड दिशानिर्देश',
        'supported_formats': 'समर्थित प्रारूप',
        'image_size': 'छवि आकार',
        'max_10mb': 'अधिकतम 10MB',
        'image_quality': 'छवि गुणवत्ता',
        'high_resolution': 'उच्च रिज़ॉल्यूशन पसंदीदा',
        'scan_type': 'स्कैन प्रकार',
        't1_t2_flair': 'T1, T2, या FLAIR अनुक्रम',
        'patient_id': 'रोगी आईडी',
        'auto_generated': 'स्वतः उत्पन्न',
        'education_years': 'शिक्षा वर्ष',
        'socioeconomic_status': 'सामाजिक-आर्थिक स्थिति',
        'cdr_score': 'CDR स्कोर',
        'cdr': 'CDR',
        'etiv': 'eTIV (अनुमानित कुल इंट्राक्रैनियल वॉल्यूम)',
        'nwbv': 'nWBV (सामान्यीकृत संपूर्ण मस्तिष्क आयतन)',
        'asf': 'ASF (एटलस स्केलिंग फैक्टर)',
        'mri_scan_upload': 'एमआरआई स्कैन अपलोड',
        'drag_drop_mri': 'अपना एमआरआई स्कैन यहाँ खींचें और छोड़ें',
        'or_click_browse': 'या फ़ाइलें ब्राउज़ करने के लिए क्लिक करें',
        'supports_formats': 'समर्थन करता है: JPEG, PNG, DICOM, NIfTI',
        'clinical_notes': 'अतिरिक्त नैदानिक नोट्स',
        'clinical_notes_placeholder': 'कोई भी प्रासंगिक नैदानिक अवलोकन या लक्षण दर्ज करें...',
        'analyze_mri_scan': 'एमआरआई स्कैन का विश्लेषण करें',
        'reset_form': 'फॉर्म रीसेट करें',
        'ai_powered_analysis': 'एआई-संचालित विश्लेषण',
        'ai_analysis_desc': 'सटीक मनोभ्रंश पहचान के लिए उन्नत मशीन लर्निंग एल्गोरिदम',
        'secure_confidential': 'सुरक्षित और गोपनीय',
        'secure_desc': 'आपका चिकित्सा डेटा एन्क्रिप्टेड है और अत्यधिक गोपनीयता के साथ संभाला जाता है',
        'fast_results': 'तेज़ परिणाम',
        'fast_results_desc': 'मिनटों में व्यापक विश्लेषण रिपोर्ट प्राप्त करें',
        'medical_disclaimer': 'चिकित्सा अस्वीकरण',
        'analyzing': 'विश्लेषण कर रहा है',
        'invalid_file_format': 'अमान्य फ़ाइल प्रारूप। कृपया JPEG, PNG, DICOM, या NIfTI फ़ाइलें अपलोड करें।',
        'file_too_large': 'फ़ाइल का आकार 10MB सीमा से अधिक है।',
        
        # Tooltips
        'patient_id_tooltip': 'आपके उपयोगकर्ता नाम और वर्तमान तिथि के आधार पर स्वचालित रूप से उत्पन्न अद्वितीय पहचानकर्ता',
        'mri_age_tooltip': 'एमआरआई स्कैन के समय रोगी की आयु। आयु महत्वपूर्ण है क्योंकि मस्तिष्क की मात्रा स्वाभाविक रूप से उम्र के साथ कम होती है।',
        'education_years_tooltip': 'पूर्ण की गई औपचारिक शिक्षा के कुल वर्ष। उच्च शिक्षा अधिक संज्ञानात्मक रिजर्व से जुड़ी है।',
        'ses_tooltip': '1-5 के पैमाने पर सामाजिक-आर्थिक स्थिति, जहां 1 सबसे अधिक है और 5 सबसे कम है। यह स्वास्थ्य देखभाल और जीवनशैली कारकों तक पहुंच को दर्शाता है।',
        'etiv_tooltip': 'अनुमानित कुल इंट्राक्रैनियल वॉल्यूम (mm³ में)। यह अधिकतम मस्तिष्क आकार को मापता है और वयस्कता भर में स्थिर रहता है। सामान्य रेंज: 1100-2000 mm³.',
        'nwbv_tooltip': 'सामान्यीकृत संपूर्ण मस्तिष्क आयतन (अनुपात)। यह इंट्राक्रैनियल वॉल्यूम के सापेक्ष मस्तिष्क की मात्रा का प्रतिनिधित्व करता है। मान आमतौर पर 0.6-0.9 से होते हैं। कम मान शोष का संकेत दे सकते हैं।',
        'asf_tooltip': 'एटलस स्केलिंग फैक्टर। व्यक्तियों के बीच मस्तिष्क के आकार की तुलना करने के लिए उपयोग किया जाने वाला सामान्यीकरण कारक। सामान्य रेंज: 0.9-1.8.',
    }
}

def get_translation(key, lang='en'):
    """Get translation for a key in specified language"""
    return TRANSLATIONS.get(lang, TRANSLATIONS['en']).get(key, key)

@app.context_processor
def inject_translations():
    """Make translation function available in all templates"""
    lang = session.get('language', 'en')
    return dict(t=lambda key: get_translation(key, lang), current_lang=lang)

@app.route('/set_language/<lang>')
def set_language(lang):
    """Set the language preference"""
    if lang in TRANSLATIONS:
        session['language'] = lang
    return redirect(request.referrer or url_for('landing'))

# === REST OF YOUR ORIGINAL CODE ===

# --- Load first model ---
try:
    model = joblib.load('alzheimers_model.pkl')
    limits = joblib.load('alzheimers_limits.pkl')
except FileNotFoundError:
    print("Alzheimer's model or limits not found. Please run train_alzheimers_model.py to generate alzheimers_model.pkl and alzheimers_limits.pkl")

# --- Load second model (MRI model) ---
try:
    mri_model = joblib.load('mri_model.pkl')
    mri_limits = joblib.load('mri_limits.pkl')
except FileNotFoundError:
    print("MRI model or limits not found. Please run train_mri_model.py to generate mri_model.pkl and mri_limits.pkl")

# --- Database Initialization ---
def init_db():
    with sqlite3.connect('users.db') as conn:
        conn.execute("""
            CREATE TABLE IF NOT EXISTS users (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                username TEXT UNIQUE,
                password TEXT,
                email TEXT
            )
        """)
        conn.execute("""
            CREATE TABLE IF NOT EXISTS predictions (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                user_id INTEGER,
                patient_name TEXT,
                prediction_result TEXT,
                input_data TEXT,
                prediction_type TEXT,
                prediction_date TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                FOREIGN KEY (user_id) REFERENCES users (id)
            )
        """)

# --- Routes ---
@app.route('/')
def landing():
    return render_template('landing.html')

@app.route('/login', methods=['GET', 'POST'])
def login():
    lang = session.get('language', 'en')
    if request.method == 'GET':
        return render_template('login.html')
    username = request.form['username']
    password = request.form['password']
    with sqlite3.connect('users.db') as conn:
        cur = conn.cursor()
        cur.execute("SELECT * FROM users WHERE username=? AND password=?", (username, password))
        user = cur.fetchone()
    if user:
        session['user'] = username
        return redirect(url_for('home'))
    else:
        return f"{get_translation('invalid_login', lang)}. <a href='/'>Back to landing</a> or <a href='/signup'>Sign up</a>."

@app.route('/signup', methods=['GET'])
def signup():
    return render_template('signup.html')

@app.route('/register', methods=['POST'])
def register():
    lang = session.get('language', 'en')
    username = request.form['new_username']
    password = request.form['new_password']
    email = request.form['email']
    with sqlite3.connect('users.db') as conn:
        try:
            conn.execute("INSERT INTO users (username, password, email) VALUES (?, ?, ?)", (username, password, email))
            return redirect(url_for('landing'))
        except sqlite3.IntegrityError:
            return f"{get_translation('username_exists', lang)}. <a href='/signup'>Try again</a> or <a href='/'>Back to landing</a>"

@app.route('/predictor')
def predictor():
    if 'user' not in session:
        return redirect(url_for('home'))
    return render_template('predictor.html', limits=limits)

@app.route('/predict', methods=['POST'])
def predict():
    lang = session.get('language', 'en')
    if 'user' not in session:
        return redirect(url_for('home'))
    patient_name = request.form.get('PatientName', 'Unknown Patient')
    input_data = {}
    for feature in ['Age','Gender','Ethnicity','EducationLevel','BMI','Smoking','AlcoholConsumption','PhysicalActivity','DietQuality','SleepQuality','FamilyHistoryAlzheimers','CardiovascularDisease','Diabetes','Depression','HeadInjury','Hypertension','SystolicBP','DiastolicBP','CholesterolTotal','CholesterolLDL','CholesterolHDL','CholesterolTriglycerides','MMSE','FunctionalAssessment','MemoryComplaints','BehavioralProblems','ADL','Confusion','Disorientation','PersonalityChanges','DifficultyCompletingTasks','Forgetfulness']:
        try: 
            value = float(request.form[feature])
            if limits[feature]['min'] <= value <= limits[feature]['max']:
                input_data[feature] = value
            else:
                return f"Value for {feature} out of range ({limits[feature]['min']} to {limits[feature]['max']}). <a href='/predictor'>Try again</a>"
        except (ValueError, KeyError):
            return f"Invalid or missing value for {feature}. <a href='/predictor'>Try again</a>"
    input_df = pd.DataFrame([input_data])
    prediction = model.predict(input_df)[0]
    result = get_translation('high_risk', lang) if prediction == 1 else get_translation('low_risk', lang)
    with sqlite3.connect('users.db') as conn:
        cur = conn.cursor()
        cur.execute("SELECT id FROM users WHERE username=?", (session['user'],))
        user_id = cur.fetchone()[0]
        input_json = json.dumps(input_data)
        cur.execute("""
            INSERT INTO predictions (user_id, patient_name, prediction_result, input_data, prediction_type)
            VALUES (?, ?, ?, ?, ?)
        """, (user_id, patient_name, result, input_json, 'clinical'))
    return render_template('results.html', result=result, patient_name=patient_name)

@app.route('/mri-upload')
def mri_upload():
    """Render MRI upload page with auto-generated patient ID"""
    if 'user' not in session:
        return redirect(url_for('login'))
    
    # Generate patient ID from username + current date
    current_date = datetime.now().strftime('%Y%m%d')  # Format: 20250127
    username = session['user']
    
    return render_template(
        'mri_upload.html', 
        mri_limits=mri_limits,
        username=username,
        current_date=current_date
    )

@app.route('/predict_mri', methods=['POST'])
def predict_mri():
    lang = session.get('language', 'en')
    if 'user' not in session:
        return redirect(url_for('home'))
    
    # Get patient name from form - use PatientID field from MRI form
    patient_name = request.form.get('PatientID', session.get('user', 'Unknown Patient'))
    
    input_data = {}
    for feature in ['Age', 'M/F', 'EDUC', 'SES', 'MMSE', 'CDR', 'eTIV', 'nWBV', 'ASF']:
        try:
            value = request.form.get(feature)
            if not value:
                return f"Missing value for {feature}. <a href='/mri-upload'>Try again</a>"
            value = float(value) if feature != 'M/F' else int(value)
            if feature == 'M/F' and value not in [0, 1]:
                return f"Invalid value for {feature} (must be 0 or 1). <a href='/mri-upload'>Try again</a>"
            if feature != 'M/F' and not (mri_limits[feature]['min'] <= value <= mri_limits[feature]['max']):
                return f"Value for {feature} out of range ({mri_limits[feature]['min']} to {mri_limits[feature]['max']}). <a href='/mri-upload'>Try again</a>"
            input_data[feature] = value
        except (ValueError, KeyError):
            return f"Invalid value for {feature}. <a href='/mri-upload'>Try again</a>"
    
    input_df = pd.DataFrame([input_data])
    prediction = mri_model.predict(input_df)[0]
    result = get_translation('high_risk', lang) if prediction == 1 else get_translation('low_risk', lang)
    
    with sqlite3.connect('users.db') as conn:
        cur = conn.cursor()
        cur.execute("SELECT id FROM users WHERE username=?", (session['user'],))
        user_id = cur.fetchone()[0]
        input_json = json.dumps(input_data)
        cur.execute("""
            INSERT INTO predictions (user_id, patient_name, prediction_result, input_data, prediction_type)
            VALUES (?, ?, ?, ?, ?)
        """, (user_id, patient_name, result, input_json, 'mri'))
    
    return render_template('results.html', result=result, patient_name=patient_name)



@app.route('/history')
def history():
    lang = session.get('language', 'en')
    if 'user' not in session:
        return f"{get_translation('must_login', lang)}. <a href='/login'>Login</a>"
    with sqlite3.connect('users.db') as conn:
        cur = conn.cursor()
        cur.execute("""
            SELECT p.patient_name, p.prediction_result, p.input_data, p.prediction_date, p.prediction_type
            FROM predictions p
            JOIN users u ON p.user_id = u.id
            WHERE u.username = ?
            ORDER BY p.prediction_date DESC
        """, (session['user'],))
        predictions = cur.fetchall()
    history_data = []
    for pred in predictions:
        patient_name, result, input_json, date, pred_type = pred
        try:
            input_data = json.loads(input_json)
            history_data.append({
                'patient_name': patient_name,
                'result': result,
                'input_data': input_data,
                'date': date,
                'type': pred_type
            })
        except:
            continue
    return render_template('history.html', predictions=history_data)

def generate_suggestions(prediction_result, input_data):
    suggestions = []
    if "HIGH RISK" in prediction_result or "அதிக ஆபத்து" in prediction_result or "उच्च जोखिम" in prediction_result:
        suggestions.append("🔴 **High Risk Detected** - Please consult with a healthcare professional immediately.")
        suggestions.append("📋 Consider scheduling a comprehensive medical evaluation.")
        suggestions.append("🧠 Monitor cognitive changes and maintain a detailed symptom diary.")
        suggestions.append("💊 Follow up with a neurologist or geriatric specialist.")
        suggestions.append("📱 Consider using memory aids and cognitive training apps.")
    else:
        suggestions.append("✅ **Low Risk Detected** - Continue with regular health monitoring.")
        suggestions.append("🧠 Maintain cognitive health through regular mental exercises.")
        suggestions.append("🏃‍♂️ Stay physically active and maintain a healthy lifestyle.")
        suggestions.append("🥗 Follow a brain-healthy diet (Mediterranean diet recommended).")
        suggestions.append("😴 Ensure adequate sleep and stress management.")
    if 'Age' in input_data and input_data['Age'] > 40:
        suggestions.append("👴 Age-related monitoring: Consider annual cognitive assessments.")
    if 'MMSE' in input_data and input_data['MMSE'] < 25:
        suggestions.append("📊 MMSE Score: Consider cognitive assessment tools and monitoring.")
    return suggestions

def create_pdf_report(input_data, prediction_result, suggestions, patient_name):
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    story = []
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=24,
        spaceAfter=30,
        alignment=1,
        textColor=colors.darkblue
    )
    story.append(Paragraph(f"Alzheimer's Disease Risk Assessment Report for {patient_name}", title_style))
    story.append(Spacer(1, 20))
    date_style = ParagraphStyle(
        'DateStyle',
        parent=styles['Normal'],
        fontSize=12,
        alignment=1
    )
    story.append(Paragraph(f"Generated on: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}", date_style))
    story.append(Spacer(1, 30))
    result_style = ParagraphStyle(
        'ResultStyle',
        parent=styles['Heading2'],
        fontSize=18,
        spaceAfter=20,
        textColor=colors.red if "HIGH RISK" in prediction_result or "அதிக ஆபத்து" in prediction_result or "उच्च जोखिम" in prediction_result else colors.green
    )
    story.append(Paragraph(f"Assessment Result: {prediction_result}", result_style))
    story.append(Spacer(1, 20))
    table_data = [['Parameter', 'Value']]
    for key, value in input_data.items():
        table_data.append([key, str(value)])
    table = Table(table_data, colWidths=[2*inch, 1.5*inch])
    table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 12),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
        ('GRID', (0, 0), (-1, -1), 1, colors.black)
    ]))
    story.append(table)
    story.append(Spacer(1, 30))
    story.append(Paragraph("Recommendations & Suggestions", styles['Heading2']))
    story.append(Spacer(1, 10))
    for i, suggestion in enumerate(suggestions, 1):
        suggestion_style = ParagraphStyle(
            'SuggestionStyle',
            parent=styles['Normal'],
            fontSize=11,
            spaceAfter=8,
            leftIndent=20
        )
        story.append(Paragraph(f"{i}. {suggestion}", suggestion_style))
    story.append(Spacer(1, 30))
    disclaimer_style = ParagraphStyle(
        'DisclaimerStyle',
        parent=styles['Normal'],
        fontSize=10,
        textColor=colors.grey,
        alignment=1
    )
    story.append(Paragraph(
        "⚠️ DISCLAIMER: This report is for informational purposes only and should not replace professional medical advice. "
        "Please consult with healthcare professionals for proper diagnosis and treatment.",
        disclaimer_style
    ))
    doc.build(story)
    buffer.seek(0)
    return buffer

def create_docx_report(input_data, prediction_result, suggestions, patient_name):
    doc = Document()
    title = doc.add_heading(f'Alzheimer\'s Disease Risk Assessment Report for {patient_name}', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_para = doc.add_paragraph(f"Generated on: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}")
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    result_para = doc.add_paragraph()
    result_para.add_run("Assessment Result: ").bold = True
    result_para.add_run(prediction_result)
    if "HIGH RISK" in prediction_result or "அதிக ஆபத்து" in prediction_result or "उच्च जोखिम" in prediction_result:
        result_para.runs[1].font.color.rgb = docx.shared.RGBColor(255, 0, 0)
    else:
        result_para.runs[1].font.color.rgb = docx.shared.RGBColor(0, 128, 0)
    doc.add_paragraph()
    doc.add_heading('Input Parameters', level=1)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Parameter'
    header_cells[1].text = 'Value'
    for key, value in input_data.items():
        row_cells = table.add_row().cells
        row_cells[0].text = key
        row_cells[1].text = str(value)
    doc.add_paragraph()
    doc.add_heading('Recommendations & Suggestions', level=1)
    for i, suggestion in enumerate(suggestions, 1):
        doc.add_paragraph(f"{i}. {suggestion}", style='List Number')
    doc.add_paragraph()
    disclaimer = doc.add_paragraph()
    disclaimer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    disclaimer_run = disclaimer.add_run(
        "⚠️ DISCLAIMER: This report is for informational purposes only and should not replace professional medical advice. "
        "Please consult with healthcare professionals for proper diagnosis and treatment."
    )
    disclaimer_run.font.size = docx.shared.Pt(10)
    disclaimer_run.font.color.rgb = docx.shared.RGBColor(128, 128, 128)
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

@app.route('/download_report/<format>')
def download_report(format):
    lang = session.get('language', 'en')
    if 'user' not in session:
        return redirect(url_for('home'))
    prediction_id = request.args.get('prediction_id', None)
    with sqlite3.connect('users.db') as conn:
        cur = conn.cursor()
        if prediction_id is not None:
            cur.execute("""
                SELECT p.patient_name, p.prediction_result, p.input_data
                FROM predictions p
                JOIN users u ON p.user_id = u.id
                WHERE u.username = ?
                ORDER BY p.prediction_date DESC
                LIMIT 1 OFFSET ?
            """, (session['user'], int(prediction_id)))
        else:
            cur.execute("""
                SELECT p.patient_name, p.prediction_result, p.input_data
                FROM predictions p
                JOIN users u ON p.user_id = u.id
                WHERE u.username = ?
                ORDER BY p.prediction_date DESC
                LIMIT 1
            """, (session['user'],))
        result = cur.fetchone()
    if not result:
        return get_translation('no_prediction', lang), 404
    patient_name, prediction_result, input_json = result
    input_data = json.loads(input_json)
    suggestions = generate_suggestions(prediction_result, input_data)
    if format.lower() == 'pdf':
        buffer = create_pdf_report(input_data, prediction_result, suggestions, patient_name)
        return send_file(
            buffer,
            as_attachment=True,
            download_name=f"alzheimer_assessment_report_{patient_name.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf",
            mimetype='application/pdf'
        )
    elif format.lower() == 'docx':
        buffer = create_docx_report(input_data, prediction_result, suggestions, patient_name)
        return send_file(
            buffer,
            as_attachment=True,
            download_name=f"alzheimer_assessment_report_{patient_name.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    else:
        return "Invalid format. Use 'pdf' or 'docx'", 400

@app.route('/home')
def home():
    if 'user' not in session:
        return redirect(url_for('login'))
    with sqlite3.connect('users.db') as conn:
        cur = conn.cursor()
        cur.execute("""
            SELECT p.prediction_result, p.input_data, p.prediction_date, p.prediction_type
            FROM predictions p
            JOIN users u ON p.user_id = u.id
            WHERE u.username = ?
            ORDER BY p.prediction_date DESC
            LIMIT 3
        """, (session['user'],))
        predictions = cur.fetchall()
    total_tests = len(predictions)
    recent_tests = []
    avg_score = 0
    last_test_days = "N/A"
    risk_level = "N/A"
    risk_change = "No data"
    if predictions:
        for pred in predictions:
            result, input_json, date, pred_type = pred
            input_data = json.loads(input_json)
            score = input_data.get('MMSE', 0) / 30 * 100 if 'MMSE' in input_data else 0
            risk = "High Risk" if "HIGH RISK" in result or "அதிக ஆபத்து" in result or "उच्च जोखिम" in result else "Low Risk"
            test_date = datetime.strptime(date, '%Y-%m-%d %H:%M:%S')
            recent_tests.append({
                'date': test_date.strftime('%b %d, %Y'),
                'risk': risk,
                'score': f"{score:.0f}%",
                'type': pred_type
            })
        scores = [json.loads(pred[1]).get('MMSE', 0) / 30 * 100 for pred in predictions]
        avg_score = sum(scores) / len(scores) if scores else 0
        risk_level = "High" if "HIGH RISK" in predictions[0][0] or "அதிக ஆபத்து" in predictions[0][0] or "उच्च जोखिम" in predictions[0][0] else "Low"
        risk_change = "Stable"
        last_test_date = datetime.strptime(predictions[0][2], '%Y-%m-%d %H:%M:%S')
        days_since = (datetime.now() - last_test_date).days
        last_test_days = f"{days_since} days"
        tests_this_month = sum(1 for pred in predictions if datetime.strptime(pred[2], '%Y-%m-%d %H:%M:%S').month == datetime.now().month)
        tests_change = f"+{tests_this_month} this month" if tests_this_month > 0 else "No tests this month"
    else:
        tests_change = "No tests this month"
    notifications = [
        {"type": "info", "message": "Your cognitive test is due for this week", "time": "2 hours ago"},
        {"type": "success", "message": "New AI analysis features available", "time": "1 day ago"},
        {"type": "warning", "message": "Schedule your next MRI scan appointment", "time": "3 days ago"}
    ]
    return render_template(
        'home.html',
        username=session['user'],
        total_tests=total_tests,
        risk_level=risk_level,
        avg_score=f"{avg_score:.0f}%",
        last_test_days=last_test_days,
        tests_change=tests_change,
        recent_tests=recent_tests,
        notifications=notifications,
        notification_count=len(notifications)
    )

@app.route('/cognitive-test')
def cognitive_test():
    return redirect(url_for('predictor'))

@app.route('/results')
def results():
    lang = session.get('language', 'en')
    if 'user' not in session:
        return redirect(url_for('login'))
    
    # Get the most recent prediction for this user
    with sqlite3.connect('users.db') as conn:
        cur = conn.cursor()
        cur.execute("""
            SELECT patient_name, prediction_result, prediction_date
            FROM predictions p
            JOIN users u ON p.user_id = u.id
            WHERE u.username = ?
            ORDER BY p.prediction_date DESC
            LIMIT 1
        """, (session['user'],))
        latest_prediction = cur.fetchone()
    
    if latest_prediction:
        patient_name, result, prediction_date = latest_prediction
        current_date = datetime.strptime(prediction_date, '%Y-%m-%d %H:%M:%S').strftime('%B %d, %Y')
    else:
        patient_name = session.get('user', 'Unknown Patient')
        result = get_translation('no_data', lang)
        current_date = datetime.now().strftime('%B %d, %Y')
    
    return render_template('results.html', 
                         result=result, 
                         patient_name=patient_name,
                         current_date=current_date)



@app.route('/profile')
def profile():
    if 'user' not in session:
        return redirect(url_for('login'))
    username = session['user']
    with sqlite3.connect('users.db') as conn:
        cur = conn.cursor()
        cur.execute("SELECT email FROM users WHERE username=?", (username,))
        row = cur.fetchone()
        email = row[0] if row else ''
    return render_template('profile.html', username=username, email=email)

@app.route('/update_profile', methods=['POST'])
def update_profile():
    if 'user' not in session:
        return redirect(url_for('login'))
    username = session['user']
    email = request.form.get('email')
    with sqlite3.connect('users.db') as conn:
        cur = conn.cursor()
        cur.execute("UPDATE users SET email=? WHERE username=?", (email, username))
    return redirect(url_for('profile'))

@app.route('/logout')
def logout():
    session.clear()
    return redirect(url_for('landing'))

if __name__ == '__main__':
    init_db()
    app.run(debug=True)